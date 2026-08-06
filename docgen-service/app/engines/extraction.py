"""Document extraction engine: plain text + vision-OCR page transcription.

THE single extraction implementation (ARCHITECTURE.md "Known POC bugs" #6):
replaces the three drifted copies in the POC (``api/main.py _run_extract``,
``subsystems/collateral_reviewer/main.py extract_document_text``,
``subsystems/policy_qa/main.py extract_document_text``), all of which mirrored
the canonical ``subsystems/document_extractor/main.py``.

Behavior (ported invariants):

- **Density heuristic**: a PDF averaging >= ``TEXT_DENSITY_THRESHOLD`` (100)
  extractable chars/page has a usable text layer; below that it is treated as
  a scanned/image-only PDF and transcribed page-by-page via a vision model.
- **Vision transcription**: each page is rendered with PyMuPDF at
  ``RASTER_SCALE`` (3x) to a base64 PNG and transcribed by an independent,
  non-streaming provider call using the frozen legacy system prompt
  (``prompts/extraction_transcription.md`` — verbatim domain IP). Pages run in
  parallel on a thread pool capped at ``max_workers`` (default 8, mirroring
  the POC's ``MAX_VISION_WORKERS``).
- **Page markers**: successful vision pages begin with ``=== PAGE N ===``
  (the model is instructed to emit the marker); text-layer pages are wrapped
  in the same marker by this module. Failed pages get the verbatim legacy
  placeholder ``=== PAGE N (TRANSCRIPTION FAILED) ===``.
- **Blank/empty transcriptions count as failures** (legacy ``_commit_page``
  behavior) — a whitespace-only page result is never silently counted as
  successful while contributing nothing downstream.
- **Keep-sparse-text-per-page fallback**: ``transcribe_pdf`` without
  ``force_vision`` keeps the (possibly sparse) embedded text of a text-layer
  PDF page-by-page instead of spending vision calls — the legacy
  ``transcribe_text_embed_pdf`` path.

Progress events: where an ``emit(type, text)`` callback is provided, this
engine emits ``"event"`` messages whose text is a compact, SINGLE-encoded
JSON string (the POC double-encoded these — see ARCHITECTURE.md Jobs + SSE):
``{"event":"init","total":T}`` first, then per page
``{"event":"page_start"|"page_complete"|"page_failed","page":N,"total":T}``.

Known POC bug fixed in this port (memory): the legacy code rasterized EVERY
page to a base64 PNG up front before transcribing, ballooning memory on large
scans. Pages are now rasterized lazily inside each worker task (each worker
opens its own short-lived PyMuPDF handle — ``fitz`` documents are not
thread-safe), so at most ``max_workers`` rendered pages exist at once.

Engines are pure logic: the LLM provider, model name, and emit callback all
arrive as arguments (no FastAPI / SQLAlchemy / app.core.config imports).
"""
from __future__ import annotations

import base64
import json
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, Callable

import docx
import fitz
import pdfplumber

if TYPE_CHECKING:  # runtime-pure: engines never import app.core.config
    from app.llm.base import LLMProvider

EmitFn = Callable[[str, str], None]

# Avg extractable chars/page at or above which a PDF's text layer is trusted
# (below → scanned/image-only → vision OCR). Legacy `text_threshold=100.0`.
TEXT_DENSITY_THRESHOLD: float = 100.0

# PyMuPDF zoom matrix for vision rasterization (legacy `scale=3.0`).
RASTER_SCALE: float = 3.0

# Cap on concurrent vision-transcription calls — keeps pages parallel without
# firing one request per page at once (avoids rate-limit failures on large
# scans). Mirrors the POC's MAX_VISION_WORKERS.
MAX_VISION_WORKERS: int = 8

_PROMPTS_DIR = Path(__file__).resolve().parent / "prompts"

# Frozen domain IP — verbatim copy of subsystems/document_extractor/prompt.md.
TRANSCRIPTION_PROMPT: str = (_PROMPTS_DIR / "extraction_transcription.md").read_text(
    encoding="utf-8"
)


@dataclass
class TranscriptionResult:
    """Extracted/transcribed document text plus per-page success accounting."""

    text: str
    pages_total: int
    pages_failed: list[int] = field(default_factory=list)


# ──────────────────────────────────────────────────────────────────────────────
# Plain text extraction (port of shared/extractor.py)
# ──────────────────────────────────────────────────────────────────────────────


def extract_text(path: Path) -> str:
    """Plain text from a ``.docx`` or text-layer PDF (no OCR).

    ``.docx`` → non-empty paragraphs plus pipe-joined table rows;
    ``.pdf`` → pdfplumber per-page text joined with blank lines.
    Raises ``ValueError`` on any other suffix.
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix == ".docx":
        return _from_docx(path)
    raise ValueError(
        f"Unsupported file type: {suffix}. Please convert .doc files to .docx before use."
    )


def _from_pdf(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n\n".join(parts)


def _from_docx(path: Path) -> str:
    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


# ──────────────────────────────────────────────────────────────────────────────
# Scanned-PDF detection (100 chars/page density heuristic, via pdfplumber)
# ──────────────────────────────────────────────────────────────────────────────


def _pdf_stats(path: Path) -> tuple[int, float]:
    """(page count, average extractable chars per page) for a PDF."""
    with pdfplumber.open(path) as pdf:
        n_pages = len(pdf.pages)
        total_chars = sum(len(page.extract_text() or "") for page in pdf.pages)
    return n_pages, total_chars / max(n_pages, 1)


def pdf_text_density(path: Path) -> float:
    """Average extractable chars/page — the scanned-vs-text routing signal."""
    return _pdf_stats(path)[1]


def is_scanned_pdf(path: Path) -> bool:
    """True when the PDF has no usable text layer (density below threshold)."""
    return pdf_text_density(path) < TEXT_DENSITY_THRESHOLD


# ──────────────────────────────────────────────────────────────────────────────
# Page-by-page transcription
# ──────────────────────────────────────────────────────────────────────────────


def _failed_placeholder(page_num: int) -> str:
    # Verbatim legacy placeholder (including the trailing newline).
    return f"=== PAGE {page_num} (TRANSCRIPTION FAILED) ===\n"


def _event(
    emit: EmitFn | None, lock: threading.Lock, payload: dict[str, int | str]
) -> None:
    """Emit one single-encoded compact-JSON "event" message (thread-safe)."""
    if emit is None:
        return
    with lock:
        emit("event", json.dumps(payload, separators=(",", ":")))


def transcribe_pdf(
    path: Path,
    provider: "LLMProvider",
    vision_model: str,
    *,
    max_workers: int = MAX_VISION_WORKERS,
    emit: EmitFn | None = None,
    force_vision: bool = False,
    prompt: str | None = None,
) -> TranscriptionResult:
    """Transcribe a PDF page-by-page into ``=== PAGE N ===``-marked text.

    Unless ``force_vision`` is set, a text-layer PDF (density >= threshold)
    keeps its embedded text per page (sparse-text fallback, no LLM calls);
    otherwise each page is rasterized lazily inside a worker task and
    transcribed by ``vision_model`` in parallel (capped at ``max_workers``).
    Failed or blank pages get placeholder text and appear in
    ``pages_failed``. Emits ``init`` / ``page_start`` / ``page_complete`` /
    ``page_failed`` events (see module docstring). ``prompt`` overrides the
    shipped transcription system prompt (None -> ``TRANSCRIPTION_PROMPT``).
    """
    if path.suffix.lower() != ".pdf":
        raise ValueError(f"transcribe_pdf requires a .pdf file, got: {path.suffix}")

    with fitz.open(str(path)) as doc:
        total = len(doc)

    lock = threading.Lock()
    _event(emit, lock, {"event": "init", "total": total})

    use_vision = force_vision or pdf_text_density(path) < TEXT_DENSITY_THRESHOLD

    pages_failed: list[int] = []
    parts: list[str] = []

    if not use_vision:
        # Text-embed PDF: keep the (possibly sparse) text layer page-by-page —
        # fast, no API calls (legacy transcribe_text_embed_pdf).
        with fitz.open(str(path)) as doc:
            for page_num in range(1, total + 1):
                _event(
                    emit, lock, {"event": "page_start", "page": page_num, "total": total}
                )
                try:
                    page_text = doc[page_num - 1].get_text()
                except Exception:
                    page_text = ""
                if page_text and page_text.strip():
                    parts.append(f"=== PAGE {page_num} ===\n\n{page_text}")
                    _event(
                        emit,
                        lock,
                        {"event": "page_complete", "page": page_num, "total": total},
                    )
                else:
                    parts.append(_failed_placeholder(page_num))
                    pages_failed.append(page_num)
                    _event(
                        emit,
                        lock,
                        {"event": "page_failed", "page": page_num, "total": total},
                    )
        return TranscriptionResult(
            text="\n\n".join(parts), pages_total=total, pages_failed=pages_failed
        )

    # Scanned PDF: parallel vision transcription, one short-lived non-streaming
    # call per page. Rasterization happens lazily INSIDE each worker (memory
    # fix — the POC rendered every page up front) using a per-task PyMuPDF
    # handle, since fitz documents are not thread-safe.
    system_prompt = prompt if prompt is not None else TRANSCRIPTION_PROMPT

    def _transcribe_page(page_num: int) -> str:
        _event(emit, lock, {"event": "page_start", "page": page_num, "total": total})
        with fitz.open(str(path)) as doc:
            pix = doc[page_num - 1].get_pixmap(
                matrix=fitz.Matrix(RASTER_SCALE, RASTER_SCALE)
            )
            img_b64 = base64.b64encode(pix.tobytes("png")).decode("utf-8")
        intro = (
            f"Transcribe page {page_num} of the document. "
            f"Begin with the === PAGE {page_num} === marker."
        )
        messages: list[dict[str, object]] = [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": intro},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{img_b64}"},
                    },
                ],
            },
        ]
        return provider.call(vision_model, messages, temperature=0.0)

    results: dict[int, str | None] = {}
    workers = max(1, min(total, max_workers))
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = {
            executor.submit(_transcribe_page, page_num): page_num
            for page_num in range(1, total + 1)
        }
        for future in as_completed(futures):
            page_num = futures[future]
            try:
                text: str | None = future.result()
            except Exception:
                text = None
            # Blank/whitespace-only transcription is a FAILURE (legacy
            # _commit_page invariant) — never counted as successful.
            if text and text.strip():
                results[page_num] = text.strip()
                _event(
                    emit,
                    lock,
                    {"event": "page_complete", "page": page_num, "total": total},
                )
            else:
                results[page_num] = None
                pages_failed.append(page_num)
                _event(
                    emit,
                    lock,
                    {"event": "page_failed", "page": page_num, "total": total},
                )

    for page_num in range(1, total + 1):
        page_text = results.get(page_num)
        parts.append(page_text if page_text else _failed_placeholder(page_num))

    return TranscriptionResult(
        text="\n\n".join(parts),
        pages_total=total,
        pages_failed=sorted(pages_failed),
    )


# ──────────────────────────────────────────────────────────────────────────────
# One-call routing entry point
# ──────────────────────────────────────────────────────────────────────────────


def extract_document(
    path: Path,
    provider: "LLMProvider",
    vision_model: str,
    *,
    emit: EmitFn | None = None,
    force_vision: bool = False,
    prompt: str | None = None,
) -> TranscriptionResult:
    """Extract a document's full text, OCR-ing scanned PDFs via vision.

    Routing: ``.docx`` and text-layer PDFs → :func:`extract_text` (no LLM,
    ``pages_failed`` empty); scanned PDFs (or ``force_vision``) →
    :func:`transcribe_pdf`. Raises ``ValueError`` on unsupported suffixes.
    ``prompt`` overrides the shipped transcription system prompt (None ->
    ``TRANSCRIPTION_PROMPT``).
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        # Legacy convention: a .docx counts as a single "page".
        return TranscriptionResult(
            text=extract_text(path), pages_total=1, pages_failed=[]
        )
    if suffix == ".pdf":
        pages_total, density = _pdf_stats(path)
        if force_vision or density < TEXT_DENSITY_THRESHOLD:
            # Routing already decided vision — pass force_vision=True so
            # transcribe_pdf skips a redundant density pass.
            return transcribe_pdf(
                path, provider, vision_model, emit=emit, force_vision=True,
                prompt=prompt,
            )
        return TranscriptionResult(
            text=extract_text(path), pages_total=pages_total, pages_failed=[]
        )
    raise ValueError(
        f"Unsupported file type: {suffix}. Please convert .doc files to .docx before use."
    )
