"""Meta-Analyzer engine: analyzes an output document template and produces a
markdown descriptor consumed by the selector and fill agents.

Ported from ``subsystems/meta_analyzer/main.py`` (POC). Pure logic: the LLM
provider, model name, template path, and domain knowledge arrive as arguments.
The frozen system prompt lives in ``prompts/meta_analyzer.md`` (verbatim copy
of the legacy ``prompt.md``).

Legacy invariants preserved:
- template text extraction matches ``shared/extractor.py`` (paragraphs +
  ``" | "``-joined table rows for .docx; page texts joined by blank lines for
  .pdf);
- the user message layout (Domain Knowledge / Document Template to Analyze /
  Filename) is byte-identical;
- ``temperature=0.0``;
- the fence strip fires only when the whole response *starts* fenced (the
  legacy ``descriptor.strip().startswith("```")`` gate), now delegating to the
  shared ``util.strip_fences``.
"""
from __future__ import annotations

from collections.abc import Callable
from functools import lru_cache
from pathlib import Path
from typing import TYPE_CHECKING, Any

import docx

from app.engines.util import strip_fences

if TYPE_CHECKING:
    from app.llm.base import LLMProvider

# emit(type, text) with type in {"reasoning", "content", "event"}
EmitFn = Callable[[str, str], None]

_PROMPT_PATH = Path(__file__).parent / "prompts" / "meta_analyzer.md"


@lru_cache(maxsize=1)
def _system_instruction() -> str:
    return _PROMPT_PATH.read_text(encoding="utf-8")


# ── template text extraction (ported verbatim from shared/extractor.py) ──────

def _extract_docx(path: Path) -> str:
    doc = docx.Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    import pdfplumber  # deferred: only needed for .pdf templates

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n\n".join(parts)


def _extract_template_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ValueError(
        f"Unsupported file type: {suffix}. "
        "Please convert .doc files to .docx before use."
    )


def _complete(
    provider: LLMProvider,
    model: str,
    messages: list[dict[str, Any]],
    emit: EmitFn | None,
) -> str:
    """call() when emit is None; stream() otherwise, forwarding each item."""
    if emit is None:
        return provider.call(model=model, messages=messages, temperature=0.0)
    parts: list[str] = []
    for item in provider.stream(model=model, messages=messages, temperature=0.0):
        kind = item.get("type", "")
        text = item.get("text", "")
        if not text:
            continue
        if kind == "content":
            parts.append(text)
        if kind in ("reasoning", "content"):
            emit(kind, text)
    return "".join(parts)


def analyze_template(
    template_path: Path,
    provider: LLMProvider,
    model: str,
    *,
    domain_knowledge: str = "",
    emit: EmitFn | None = None,
    prompt: str | None = None,
) -> str:
    """Analyze one output template; return the descriptor markdown.

    ``prompt`` overrides the shipped system prompt (None -> the frozen
    ``prompts/meta_analyzer.md``).
    """
    doc_text = _extract_template_text(template_path)
    user_message = (
        f"## Domain Knowledge\n\n{domain_knowledge}\n\n"
        f"---\n\n"
        f"## Document Template to Analyze\n\n"
        f"Filename: {template_path.name}\n\n"
        f"{doc_text}"
    )
    messages: list[dict[str, Any]] = [
        {"role": "system", "content": prompt if prompt is not None else _system_instruction()},
        {"role": "user", "content": user_message},
    ]
    descriptor = _complete(provider, model, messages, emit)

    if descriptor.strip().startswith("```"):
        descriptor = strip_fences(descriptor)
    return descriptor
