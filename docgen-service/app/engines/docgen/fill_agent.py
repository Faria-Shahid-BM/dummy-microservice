"""Fill Agent engine: populates ONE output template instance from case data.

Ported from ``subsystems/fill_agent/main.py`` (POC). Pure logic: template
path, descriptor, selector evidence, entity scope, case text, domain
knowledge, provider and model all arrive as arguments; the caller saves the
returned document and provenance. The frozen system prompt lives in
``prompts/fill_agent.md`` (verbatim copy of the legacy ``prompt.md``).

CRITICAL INVARIANT: the SAME python-docx ``Document`` object that
``serialize_template`` enumerated is the one ``execute_operations`` mutates —
the ``[para:i]`` / ``[cell:t,r,c]`` coordinates the model returns are valid by
construction only against that object. ``fill_document`` guarantees this.

The run-aware matchers (``BLANK_CHARS``, ``_flexible_pattern``, ``_find_span``,
``_replace_in_paragraph``, ``_set_cell_value``) are ported byte-for-byte —
they are the hardest-won IP in the POC.

POC bugs fixed here (ARCHITECTURE.md "Known POC bugs"):
- #2: ``execute_operations`` caught only ``IndexError``/``KeyError``, so a
  type-confused op aborted the whole task — it now catches ALL exceptions
  per-operation, records ``op["_error"]`` and continues; a no-match
  ``replace_text`` also gets a human-readable ``_error`` note.
- #3 (shared): response parsing goes through the tolerant
  ``util.parse_json_response`` (``strict=False`` inside, so multi-line string
  values with raw newlines survive; python-docx's text setter turns ``\\n``
  into ``<w:br/>`` on write). ``EngineParseError`` propagates with the raw
  response attached — the caller persists it.
"""
from __future__ import annotations

import json
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from functools import lru_cache
from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from app.engines.util import parse_json_response

if TYPE_CHECKING:
    from app.llm.base import LLMProvider

# emit(type, text) with type in {"reasoning", "content", "event"}
EmitFn = Callable[[str, str], None]

_PROMPT_PATH = Path(__file__).parent / "prompts" / "fill_agent.md"


@lru_cache(maxsize=1)
def _system_instruction() -> str:
    return _PROMPT_PATH.read_text(encoding="utf-8")


# ── template serialization ────────────────────────────────────────────────────

def serialize_template(doc: DocumentObject) -> str:
    lines = ["=== TOP-LEVEL PARAGRAPHS ==="]
    for i, para in enumerate(doc.paragraphs):
        lines.append(f"[para:{i}] {para.text}")
    lines.append("\n=== TABLES ===")
    for t, table in enumerate(doc.tables):
        lines.append(f"[table:{t}]")
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                lines.append(f"  [cell:{t},{r},{c}] {cell.text.strip()}")
    return "\n".join(lines)


# ── run-aware text replacement ────────────────────────────────────────────────

BLANK_CHARS = "._…\u2002\u2003 \t-"


def _flexible_pattern(find: str) -> str:
    pattern = []
    i = 0
    while i < len(find):
        ch = find[i]
        if ch in BLANK_CHARS:
            j = i
            while j < len(find) and find[j] in BLANK_CHARS:
                j += 1
            pattern.append(r"[\._…\u2002\u2003 \t\-]+")
            i = j
        else:
            pattern.append(re.escape(ch))
            i += 1
    return "".join(pattern)


def _find_span(full: str, find: str, occurrence: int = 1) -> tuple[int, int] | None:
    start, search_from = -1, 0
    for _ in range(occurrence):
        start = full.find(find, search_from)
        if start == -1:
            break
        search_from = start + len(find)
    if start != -1:
        return start, start + len(find)
    try:
        matches = list(re.finditer(_flexible_pattern(find), full))
    except re.error:
        return None
    if len(matches) >= occurrence:
        m = matches[occurrence - 1]
        return m.start(), m.end()
    return None


def _replace_in_paragraph(
    para: Paragraph, find: str, replace: str, occurrence: int = 1
) -> bool:
    runs = para.runs
    if not runs:
        return False
    texts = [r.text for r in runs]
    full = "".join(texts)
    span = _find_span(full, find, occurrence)
    if span is None:
        return False
    start, end = span
    run_starts, pos = [], 0
    for t in texts:
        run_starts.append(pos)
        pos += len(t)

    def locate(idx: int) -> tuple[int, int]:
        for i in range(len(runs)):
            if run_starts[i] <= idx < run_starts[i] + len(texts[i]):
                return i, idx - run_starts[i]
        return len(runs) - 1, len(texts[-1])

    s_run, s_off = locate(start)
    e_run, e_off = locate(end)
    if s_run == e_run:
        t = texts[s_run]
        runs[s_run].text = t[:s_off] + replace + t[e_off:]
    else:
        runs[s_run].text = texts[s_run][:s_off] + replace
        for i in range(s_run + 1, e_run):
            runs[i].text = ""
        runs[e_run].text = texts[e_run][e_off:]
    return True


def _set_cell_value(cell: _Cell, value: str) -> None:
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = value
        for r in first.runs[1:]:
            r.text = ""
    else:
        first.add_run(value)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def _no_match_note(op: dict[str, Any]) -> str:
    """Human-readable note for a replace_text op whose find text never matched."""
    find_preview = str(op.get("find", ""))
    if len(find_preview) > 80:
        find_preview = find_preview[:77] + "..."
    return f"replace_text: no match for {find_preview!r} at the target location"


def execute_operations(
    doc: DocumentObject, operations: list[dict[str, Any]]
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Apply fill operations to ``doc`` (mutating it); return (applied, failed).

    Ops are annotated in place: a failed op carries ``op["_error"]``.
    POC bug #2 fixed: ALL exceptions are caught per-operation (recorded on the
    op, execution continues), not just IndexError/KeyError.
    """
    applied, failed = [], []
    for op in operations:
        op_type = op.get("type")
        ok = False
        try:
            if op_type == "set_cell":
                cell = doc.tables[op["table_index"]] \
                    .rows[op["row_index"]].cells[op["col_index"]]
                _set_cell_value(cell, op["value"])
                ok = True
            elif op_type == "replace_text":
                find = op["find"]
                replace = op["replace"]
                occ = op.get("occurrence", 1)
                if "para_index" in op and op.get("para_index") is not None:
                    para = doc.paragraphs[op["para_index"]]
                    ok = _replace_in_paragraph(para, find, replace, occ)
                elif "table_index" in op:
                    cell = doc.tables[op["table_index"]] \
                        .rows[op["row_index"]].cells[op["col_index"]]
                    for para in cell.paragraphs:
                        if _replace_in_paragraph(para, find, replace, occ):
                            ok = True
                            break
        except Exception as exc:  # POC bug #2 fix: never abort the batch
            op["_error"] = str(exc)
        if not ok and op_type == "replace_text" and "_error" not in op:
            op["_error"] = _no_match_note(op)
        (applied if ok else failed).append(op)
    return applied, failed


# ── prompt assembly ───────────────────────────────────────────────────────────

def build_messages(
    system_instruction: str,
    domain_knowledge: str,
    descriptor: str,
    selector_evidence: str,
    entity_scope: str | None,
    template_text: str,
    case_text: str,
) -> list[dict[str, Any]]:
    text = (
        f"## Domain Knowledge\n\n{domain_knowledge}\n\n---\n\n"
        f"## Descriptor For This Template\n\n{descriptor}\n\n---\n\n"
        f"## Why This Document Was Selected\n\n{selector_evidence}\n\n"
    )
    if entity_scope:
        text += (
            f"\n**This instance is specifically for:** {entity_scope}\n"
            f"Fill ONLY the data relevant to this entity.\n\n"
        )
    text += (
        f"---\n\n## Output Template (to be filled)\n\n"
        f"The template is serialized below with location markers. "
        f"Produce fill operations referencing these locations.\n\n"
        f"{template_text}\n\n---\n\n"
        f"## Input Document (transcribed text)\n\n"
        f"All values you fill must come from this text.\n\n{case_text}"
    )
    return [
        {"role": "system", "content": system_instruction},
        {"role": "user", "content": text},
    ]


# ── high-level fill ───────────────────────────────────────────────────────────

@dataclass
class FillResult:
    """Outcome of one template-instance fill; the caller saves ``document``."""

    document: DocumentObject
    applied: list[dict[str, Any]] = field(default_factory=list)
    failed: list[dict[str, Any]] = field(default_factory=list)
    unfilled_fields: list[Any] = field(default_factory=list)


def _complete(
    provider: LLMProvider,
    model: str,
    messages: list[dict[str, Any]],
    emit: EmitFn | None,
) -> str:
    """call() when emit is None; stream() otherwise, forwarding each item."""
    if emit is None:
        return provider.call(model=model, messages=messages, temperature=0.0)
    parts: list[str] = []
    for item in provider.stream(model=model, messages=messages, temperature=0.0):
        kind = item.get("type", "")
        text = item.get("text", "")
        if not text:
            continue
        if kind == "content":
            parts.append(text)
        if kind in ("reasoning", "content"):
            emit(kind, text)
    return "".join(parts)


def _event(emit: EmitFn, payload: dict[str, Any]) -> None:
    """Emit a compact single-encoded JSON event (the POC double-encoded these)."""
    emit("event", json.dumps(payload, ensure_ascii=False, separators=(",", ":")))


def fill_document(
    template_path: Path,
    descriptor: str,
    evidence: str,
    entity_scope: str | None,
    case_text: str,
    provider: LLMProvider,
    model: str,
    *,
    domain_knowledge: str = "",
    emit: EmitFn | None = None,
    prompt: str | None = None,
) -> FillResult:
    """Fill one template instance; return the mutated document + provenance.

    Opens the template, serializes THAT document object for the model, applies
    the returned operations to THAT SAME object (coordinates valid by
    construction), and returns it unsaved — the caller persists the .docx and
    the applied/failed/unfilled provenance.

    Raises ``EngineParseError`` (raw response attached) when the model output
    cannot be parsed as JSON — the caller persists the raw text.

    ``prompt`` overrides the shipped system prompt (None -> the frozen
    ``prompts/fill_agent.md``).
    """
    doc = Document(str(template_path))
    template_text = serialize_template(doc)

    messages = build_messages(
        prompt if prompt is not None else _system_instruction(),
        domain_knowledge, descriptor,
        evidence, entity_scope, template_text, case_text,
    )
    response = _complete(provider, model, messages, emit)

    result = parse_json_response(response)
    operations: list[dict[str, Any]] = result.get("operations", [])
    unfilled_fields: list[Any] = result.get("unfilled_fields", [])

    if emit is not None:
        _event(emit, {"stage": "operations", "count": len(operations)})

    applied, failed = execute_operations(doc, operations)

    if emit is not None:
        for status, ops in (("applied", applied), ("failed", failed)):
            for op in ops:
                payload: dict[str, Any] = {
                    "stage": "op",
                    "status": status,
                    "type": op.get("type"),
                    "field": op.get("field"),
                }
                if op.get("_error"):
                    payload["error"] = op["_error"]
                _event(emit, payload)
        _event(emit, {
            "stage": "fill_done",
            "applied": len(applied),
            "failed": len(failed),
            "unfilled": len(unfilled_fields),
        })

    return FillResult(
        document=doc,
        applied=applied,
        failed=failed,
        unfilled_fields=unfilled_fields,
    )
