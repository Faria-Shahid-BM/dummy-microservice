from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from pathlib import Path
import difflib
import os
import re
import tempfile

import mammoth
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document

from claims import get_claims
from audit import audit

app = FastAPI()

ALLOWED_SUFFIXES = {".docx", ".pdf"}
TEXT_DENSITY_THRESHOLD = 100.0
_TOKEN_RE = re.compile(r"\S+|\s+")
_HTML_TOKEN_RE = re.compile(r"<[^>]+>|[^\s<]+|\s+")
_WORD_RE = re.compile(r"\S+")

MISSING_PAGE_COVERAGE_THRESHOLD = 0.2
MISSING_PAGE_MIN_MATCH_WORDS = 3
LARGE_SECTION_WORD_THRESHOLD = 40


def _is_large_removal(change_type: str, before_text: str) -> bool:
    return change_type == "deletion" and len(_WORD_RE.findall(before_text)) >= LARGE_SECTION_WORD_THRESHOLD


# --- plain text extraction (used for .pdf, and as a mixed-type fallback) ---

def _pdf_pages(path: Path) -> list[str]:
    with pdfplumber.open(path) as pdf:
        return [page.extract_text() or "" for page in pdf.pages]


def _from_pdf(path: Path) -> str:
    return "\n\n".join(_pdf_pages(path))


def find_missing_pages(pages: list[str], returned_text: str) -> list[int]:
    """Page numbers (1-indexed) whose content doesn't meaningfully appear in returned_text."""
    returned_words = _WORD_RE.findall(returned_text)
    missing = []
    for page_num, page_text in enumerate(pages, start=1):
        page_words = _WORD_RE.findall(page_text)
        if not page_words:
            continue
        matcher = difflib.SequenceMatcher(a=page_words, b=returned_words, autojunk=False)
        matched = sum(b.size for b in matcher.get_matching_blocks() if b.size >= MISSING_PAGE_MIN_MATCH_WORDS)
        if matched / len(page_words) < MISSING_PAGE_COVERAGE_THRESHOLD:
            missing.append(page_num)
    return missing


def _from_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix == ".docx":
        return _from_docx(path)
    raise ValueError(f"unsupported file type: {suffix}")


def is_scanned_pdf(path: Path) -> bool:
    if path.suffix.lower() != ".pdf":
        return False
    pages = _pdf_pages(path)
    if not pages:
        return False
    total_chars = sum(len(p) for p in pages)
    return (total_chars / len(pages)) < TEXT_DENSITY_THRESHOLD


# --- plain-text word diff (pdf, and mixed-type fallback) -------------------------------------------------------
# Deterministic word-level redline — stdlib only (difflib + re). Each
# meaningful change gets a sequential `id`; the delete/insert segment(s) that
# produced it carry the same id, which is what the frontend scrolls to when
# you click a row in the changes table.

def _tokenize(text: str) -> list[str]:
    return _TOKEN_RE.findall(text)


def _normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def compare_documents(original_text: str, returned_text: str) -> dict:
    original_text = _normalize_newlines(original_text)
    returned_text = _normalize_newlines(returned_text)

    a_tokens = _tokenize(original_text)
    b_tokens = _tokenize(returned_text)

    matcher = difflib.SequenceMatcher(a=a_tokens, b=b_tokens, autojunk=False)
    similarity = matcher.ratio()

    segments = []
    changes = []
    counts = {"insertions": 0, "deletions": 0, "replacements": 0}
    type_to_count_key = {"deletion": "deletions", "insertion": "insertions", "replacement": "replacements"}
    tag_to_type = {"delete": "deletion", "insert": "insertion", "replace": "replacement"}
    next_id = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            segments.append({"op": "equal", "text": "".join(a_tokens[i1:i2]), "changeId": None})
            continue

        before = "".join(a_tokens[i1:i2]) if tag in ("delete", "replace") else ""
        after = "".join(b_tokens[j1:j2]) if tag in ("insert", "replace") else ""

        is_meaningful = not (before.strip() == "" and after.strip() == "")
        this_id = next_id if is_meaningful else None

        if before:
            segments.append({"op": "delete", "text": before, "changeId": this_id})
        if after:
            segments.append({"op": "insert", "text": after, "changeId": this_id})

        if not is_meaningful:
            continue

        change_type = tag_to_type[tag]
        changes.append({
            "id": this_id,
            "type": change_type,
            "before": before,
            "after": after,
            "possibleMissingSection": _is_large_removal(change_type, before),
        })
        counts[type_to_count_key[change_type]] += 1
        next_id += 1

    return {
        "render": "text",
        "identical": len(changes) == 0,
        "similarity": similarity,
        "summary": {**counts, "changes": len(changes)},
        "changes": changes,
        "segments": segments,
    }


# --- docx -> real HTML, diffed in place -------------------------------------------------------
# Converts both files to actual formatted HTML (real tables, headings, bold —
# via mammoth) instead of flattening to plain text, then runs the SAME kind
# of token diff as above, except tokens are either "a whole HTML tag" or "a
# word/whitespace run" — tags flow through untouched, only word tokens get
# wrapped in <del>/<ins>. Because both documents usually come from the same
# template, the vast majority of tag tokens line up as equal and only the
# actual changed words get marked. Any tag imbalance the token diff
# introduces (rare, and more likely on documents with real structural
# changes rather than value edits) gets auto-repaired by the BeautifulSoup
# re-parse at the end — same fix-up behavior a browser applies to slightly
# malformed HTML.

def _docx_to_html(path: Path) -> str:
    with open(path, "rb") as f:
        return mammoth.convert_to_html(f).value


def _tokenize_html(html: str) -> list[str]:
    return _HTML_TOKEN_RE.findall(html)


def _strip_html(fragment: str) -> str:
    text = BeautifulSoup(fragment, "html.parser").get_text(separator=" ")
    return re.sub(r"\s+", " ", text).strip()


def compare_documents_html(original_html: str, returned_html: str) -> dict:
    a_tokens = _tokenize_html(original_html)
    b_tokens = _tokenize_html(returned_html)

    matcher = difflib.SequenceMatcher(a=a_tokens, b=b_tokens, autojunk=False)
    similarity = matcher.ratio()

    out = []
    changes = []
    counts = {"insertions": 0, "deletions": 0, "replacements": 0}
    type_to_count_key = {"deletion": "deletions", "insertion": "insertions", "replacement": "replacements"}
    tag_to_type = {"delete": "deletion", "insert": "insertion", "replace": "replacement"}
    next_id = 0

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            out.append("".join(a_tokens[i1:i2]))
            continue

        before = "".join(a_tokens[i1:i2]) if tag in ("delete", "replace") else ""
        after = "".join(b_tokens[j1:j2]) if tag in ("insert", "replace") else ""

        is_meaningful = not (before.strip() == "" and after.strip() == "")
        this_id = next_id if is_meaningful else None
        # Angular's [innerHTML] sanitizer strips `id` (it's not on the
        # attribute allowlist) but keeps `class` verbatim, so the anchor for
        # click-to-jump has to ride along as a second class instead of an id.
        anchor_class = f" doc-change-{this_id}" if this_id is not None else ""

        if before.strip():
            out.append(f'<del class="seg-delete{anchor_class}">{before}</del>')
        elif before:
            out.append(before)  # pure whitespace/tag noise — don't wrap visibly

        if after.strip():
            out.append(f'<ins class="seg-insert{anchor_class}">{after}</ins>')
        elif after:
            out.append(after)

        if not is_meaningful:
            continue

        change_type = tag_to_type[tag]
        before_text = _strip_html(before)
        changes.append({
            "id": this_id,
            "type": change_type,
            "before": before_text,
            "after": _strip_html(after),
            "possibleMissingSection": _is_large_removal(change_type, before_text),
        })
        counts[type_to_count_key[change_type]] += 1
        next_id += 1

    merged = BeautifulSoup("".join(out), "html.parser")

    return {
        "render": "html",
        "identical": len(changes) == 0,
        "similarity": similarity,
        "summary": {**counts, "changes": len(changes)},
        "changes": changes,
        "html": str(merged),
    }


# --- HTTP layer -------------------------------------------------------

def _save_upload(upload: UploadFile) -> Path:
    suffix = Path(upload.filename or "").suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise HTTPException(status_code=400, detail=f"unsupported file type: {suffix or '(none)'}")
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "wb") as f:
        f.write(upload.file.read())
    return Path(tmp_path)


def _check_not_scanned(path: Path, filename: str | None) -> None:
    if is_scanned_pdf(path):
        raise HTTPException(status_code=400, detail=f"{filename} looks like a scanned PDF with no extractable text")


@app.post("/compare")
def compare(
    original: UploadFile = File(...),
    returned: UploadFile = File(...),
    claims: dict = Depends(get_claims),
):
    original_suffix = Path(original.filename or "").suffix.lower()
    returned_suffix = Path(returned.filename or "").suffix.lower()

    original_path = _save_upload(original)
    returned_path = _save_upload(returned)
    try:
        _check_not_scanned(original_path, original.filename)
        _check_not_scanned(returned_path, returned.filename)

        if original_suffix == ".docx" and returned_suffix == ".docx":
            original_html = _docx_to_html(original_path)
            returned_html = _docx_to_html(returned_path)
            if not BeautifulSoup(original_html, "html.parser").get_text().strip():
                raise HTTPException(status_code=400, detail=f"{original.filename} has no extractable text")
            if not BeautifulSoup(returned_html, "html.parser").get_text().strip():
                raise HTTPException(status_code=400, detail=f"{returned.filename} has no extractable text")
            result = compare_documents_html(original_html, returned_html)
        else:
            original_pages = _pdf_pages(original_path) if original_suffix == ".pdf" else None
            original_text = "\n\n".join(original_pages) if original_pages is not None else extract_text(original_path)
            returned_text = extract_text(returned_path)
            if not original_text.strip():
                raise HTTPException(status_code=400, detail=f"{original.filename} has no extractable text")
            if not returned_text.strip():
                raise HTTPException(status_code=400, detail=f"{returned.filename} has no extractable text")
            result = compare_documents(original_text, returned_text)
            if original_pages is not None:
                result["missingPages"] = find_missing_pages(original_pages, returned_text)
    finally:
        original_path.unlink(missing_ok=True)
        returned_path.unlink(missing_ok=True)

    user = claims.get("username", "anonymous")
    audit(user, "document-reviewer", "compare.called")
    return result


@app.get("/health-doc")
def health():
    return {"status": "ok", "service": "document-reviewer"}
