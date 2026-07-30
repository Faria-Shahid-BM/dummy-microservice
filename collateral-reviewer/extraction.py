"""Document extraction: plain text + vision-OCR page transcription for scanned PDFs.

Ported from cad-workbench's app/engines/extraction.py, trimmed to what
collateral review actually calls (extract_document, is_scanned_pdf). The LLM
provider, model name, and emit callback all arrive as arguments — this module
has no config/settings of its own.
"""
from __future__ import annotations

import base64
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable

import docx
import fitz
import pdfplumber

EmitFn = Callable[[str, str], None]

# Avg extractable chars/page at or above which a PDF's text layer is trusted;
# below that it's treated as scanned/image-only and sent to vision OCR.
TEXT_DENSITY_THRESHOLD: float = 100.0

# PyMuPDF zoom matrix for vision rasterization.
RASTER_SCALE: float = 3.0

# Cap on concurrent vision-transcription calls.
MAX_VISION_WORKERS: int = 8

_PROMPTS_DIR = Path(__file__).resolve().parent / "prompts"
TRANSCRIPTION_PROMPT: str = (_PROMPTS_DIR / "extraction_transcription.md").read_text(
    encoding="utf-8"
)


@dataclass
class TranscriptionResult:
    text: str
    pages_total: int
    pages_failed: list[int] = field(default_factory=list)


# --- plain text extraction (no OCR) -------------------------------------------------------

def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix == ".docx":
        return _from_docx(path)
    raise ValueError(f"unsupported file type: {suffix}")


def _from_pdf(path: Path) -> str:
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n\n".join(parts)


def _from_docx(path: Path) -> str:
    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


# --- scanned-PDF detection -------------------------------------------------------

def _pdf_stats(path: Path) -> tuple[int, float]:
    with pdfplumber.open(path) as pdf:
        n_pages = len(pdf.pages)
        total_chars = sum(len(page.extract_text() or "") for page in pdf.pages)
    return n_pages, total_chars / max(n_pages, 1)


def pdf_text_density(path: Path) -> float:
    return _pdf_stats(path)[1]


def is_scanned_pdf(path: Path) -> bool:
    if path.suffix.lower() != ".pdf":
        return False
    return pdf_text_density(path) < TEXT_DENSITY_THRESHOLD


# --- page-by-page vision transcription -------------------------------------------------------

def _failed_placeholder(page_num: int) -> str:
    return f"=== PAGE {page_num} (TRANSCRIPTION FAILED) ===\n"


def transcribe_pdf(
    path: Path,
    provider,
    vision_model: str,
    *,
    max_workers: int = MAX_VISION_WORKERS,
    force_vision: bool = False,
    prompt: str | None = None,
) -> TranscriptionResult:
    if path.suffix.lower() != ".pdf":
        raise ValueError(f"transcribe_pdf requires a .pdf file, got: {path.suffix}")

    with fitz.open(str(path)) as doc:
        total = len(doc)

    use_vision = force_vision or pdf_text_density(path) < TEXT_DENSITY_THRESHOLD
    pages_failed: list[int] = []
    parts: list[str] = []

    if not use_vision:
        # text-embed PDF: keep the (possibly sparse) text layer, no API calls
        with fitz.open(str(path)) as doc:
            for page_num in range(1, total + 1):
                try:
                    page_text = doc[page_num - 1].get_text()
                except Exception:
                    page_text = ""
                if page_text and page_text.strip():
                    parts.append(f"=== PAGE {page_num} ===\n\n{page_text}")
                else:
                    parts.append(_failed_placeholder(page_num))
                    pages_failed.append(page_num)
        return TranscriptionResult(text="\n\n".join(parts), pages_total=total, pages_failed=pages_failed)

    # scanned PDF: parallel vision transcription, rasterizing lazily inside
    # each worker (fitz documents aren't thread-safe, so each task opens its
    # own short-lived handle instead of pre-rendering every page up front).
    system_prompt = prompt if prompt is not None else TRANSCRIPTION_PROMPT

    def _transcribe_page(page_num: int) -> str | None:
        with fitz.open(str(path)) as doc:
            pix = doc[page_num - 1].get_pixmap(matrix=fitz.Matrix(RASTER_SCALE, RASTER_SCALE))
            img_b64 = base64.b64encode(pix.tobytes("png")).decode("utf-8")
        intro = (
            f"Transcribe page {page_num} of the document. "
            f"Begin with the === PAGE {page_num} === marker."
        )
        messages = [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": intro},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}},
                ],
            },
        ]
        return provider.call(vision_model, messages, temperature=0.0)

    results: dict[int, str | None] = {}
    workers = max(1, min(total, max_workers))
    with ThreadPoolExecutor(max_workers=workers) as executor:
        futures = {executor.submit(_transcribe_page, page_num): page_num for page_num in range(1, total + 1)}
        for future in as_completed(futures):
            page_num = futures[future]
            try:
                text = future.result()
            except Exception:
                text = None
            # blank/whitespace-only transcription counts as a failure
            if text and text.strip():
                results[page_num] = text.strip()
            else:
                results[page_num] = None
                pages_failed.append(page_num)

    for page_num in range(1, total + 1):
        page_text = results.get(page_num)
        parts.append(page_text if page_text else _failed_placeholder(page_num))

    return TranscriptionResult(text="\n\n".join(parts), pages_total=total, pages_failed=sorted(pages_failed))


# --- one-call routing entry point -------------------------------------------------------

def extract_document(
    path: Path,
    provider,
    vision_model: str,
    *,
    force_vision: bool = False,
    prompt: str | None = None,
) -> TranscriptionResult:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return TranscriptionResult(text=extract_text(path), pages_total=1, pages_failed=[])
    if suffix == ".pdf":
        pages_total, density = _pdf_stats(path)
        if force_vision or density < TEXT_DENSITY_THRESHOLD:
            return transcribe_pdf(path, provider, vision_model, force_vision=True, prompt=prompt)
        return TranscriptionResult(text=extract_text(path), pages_total=pages_total, pages_failed=[])
    raise ValueError(f"unsupported file type: {suffix}")
